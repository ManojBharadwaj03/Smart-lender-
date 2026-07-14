import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(PROJECT_ROOT, "Smart_Lender_Project.docx")
FILES_TO_INCLUDE = [
    "README.md",
    "requirements.txt",
    "train_model.py",
    "app.py",
    ".gitignore",
    os.path.join("data", "loan_data.csv"),
    os.path.join("templates", "index.html"),
    os.path.join("templates", "result.html"),
]


def read_file_contents(path):
    absolute_path = os.path.join(PROJECT_ROOT, path)
    if not os.path.exists(absolute_path):
        return None
    with open(absolute_path, "r", encoding="utf-8", errors="replace") as handle:
        return handle.read()


def read_file_preview(path, max_lines=14):
    absolute_path = os.path.join(PROJECT_ROOT, path)
    if not os.path.exists(absolute_path):
        return None
    with open(absolute_path, "r", encoding="utf-8", errors="replace") as handle:
        lines = [line.rstrip() for _, line in zip(range(max_lines), handle)]
    return "\n".join(lines)


def add_heading(document, text, level=1):
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if level == 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
    return heading


def add_paragraph(document, text, bold=False):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_code_block(document, text):
    for line in text.splitlines():
        p = document.add_paragraph(style="Code")
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def add_full_file_section(document, path):
    document.add_heading(f"File: {path}", level=3)
    content = read_file_contents(path)
    if content is None:
        document.add_paragraph("File not found.")
        return
    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run("Full file contents:")
    title_run.bold = True
    add_code_block(document, content)


def build_document():
    document = Document()
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(11)
    if "Code" not in document.styles:
        style = document.styles.add_style("Code", 1)
        style.font.name = "Courier New"
        style.font.size = Pt(9)

    add_heading(document, "Smart Lender - Loan Eligibility Prediction", level=1)
    add_paragraph(document, f"Generated on {datetime.now():%Y-%m-%d %H:%M:%S}")
    document.add_paragraph("Smart Lender is a machine learning web application that predicts loan eligibility using applicant data, multiple classification models, and a Flask web interface.")
    document.add_paragraph("This document contains a complete project summary, setup instructions, dataset description, model pipeline details, and full source file contents.")

    document.add_heading("Contents", level=2)
    for item in ["Project Overview", "Key Features", "Environment Setup", "Dataset and Columns", "Model Training", "Web App Structure", "Project Files", "Full Project Source Contents"]:
        document.add_paragraph(item, style="List Number")

    document.add_heading("Project Overview", level=2)
    readme_content = read_file_contents("README.md")
    if readme_content:
        for paragraph in readme_content.split("\n\n"):
            document.add_paragraph(paragraph.replace("\n", " "))
    else:
        document.add_paragraph("README file not found.")

    document.add_heading("Key Features", level=2)
    for feature in [
        "Preprocess loan applicant data including categorical encoding and numeric scaling.",
        "Train and compare Decision Tree, Random Forest, KNN, and XGBoost models.",
        "Select and save the best-performing model for deployment.",
        "Provide a Flask web interface for real-time loan eligibility prediction.",
    ]:
        document.add_paragraph(feature, style="List Bullet")

    document.add_heading("Environment Setup", level=2)
    for step in [
        "Create and activate a Python virtual environment.",
        "Install required packages using requirements.txt.",
        "Train the machine learning model before launching the app.",
        "Run the Flask app and open the local web interface to predict loan eligibility.",
    ]:
        document.add_paragraph(step, style="List Bullet")
    document.add_paragraph("Recommended commands:")
    add_code_block(document, "python -m venv venv\nvenv\\Scripts\\activate\npip install -r requirements.txt\npython train_model.py\npython app.py")

    document.add_heading("Dataset and Columns", level=2)
    document.add_paragraph("The loan dataset includes applicant demographic and financial fields used to predict loan status. The main columns are:", style="Normal")
    for column in ["Gender", "Married", "Education", "Self_Employed", "ApplicantIncome", "CoapplicantIncome", "LoanAmount", "Loan_Amount_Term", "Credit_History", "Property_Area", "Loan_Status"]:
        document.add_paragraph(column, style="List Bullet")

    document.add_heading("Model Training", level=2)
    document.add_paragraph("The training pipeline reads the dataset, fills missing values, encodes categorical features, scales numeric fields, and trains multiple classifiers. The best model is selected based on test accuracy.", style="Normal")
    document.add_paragraph("Included classification models:", style="Normal")
    for model in ["Decision Tree", "Random Forest", "K-Nearest Neighbors (KNN)", "XGBoost"]:
        document.add_paragraph(model, style="List Bullet")

    document.add_heading("Web App Structure", level=2)
    document.add_paragraph("The Flask web app loads the saved model and renders an input form. Users submit applicant details and receive an approval prediction with confidence.", style="Normal")
    document.add_paragraph("Key app behavior:", style="Normal")
    for item in [
        "Render the home page with an applicant input form.",
        "Collect user inputs and convert them into model features.",
        "Predict loan approval and display the result.",
        "Show the confidence score when probability data is available.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Project Files", level=2)
    for path in FILES_TO_INCLUDE:
        document.add_paragraph(path, style="List Bullet")

    document.add_heading("Full Project Source Contents", level=2)
    for path in FILES_TO_INCLUDE:
        add_full_file_section(document, path)

    return document


def main():
    doc = build_document()
    doc.save(OUTPUT_FILE)
    print(f"DOCX generated: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
